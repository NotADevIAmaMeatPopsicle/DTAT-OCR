"""
Extraction Pipeline with Retry Logic and Escalation

Tries extraction methods from cheapest to most expensive:
1. Native extraction (free)
2. Local OCR - LightOnOCR (your compute)
3. AWS Textract (paid, optional)
4. Dead Letter Queue (manual review)
"""

import re
import time
from datetime import datetime
from pathlib import Path
from typing import Optional, Tuple
from dataclasses import dataclass, asdict
import traceback

from config import config, ProcessingConfig
from database import (
    DocumentRecord, ProcessingStatus, ExtractionMethod,
    update_document, log_processing_attempt
)


# ==============================================================================
# Normalized Data Structures (Multi-Format Output Support)
# ==============================================================================

@dataclass
class Point:
    """2D point with normalized coordinates (0.0-1.0)"""
    x: float
    y: float

    def to_dict(self):
        return {"x": self.x, "y": self.y}

    @classmethod
    def from_dict(cls, data: dict):
        return cls(x=data["x"], y=data["y"])


@dataclass
class BoundingBox:
    """Normalized bounding box (0.0-1.0 coordinates)"""
    left: float
    top: float
    width: float
    height: float

    def to_dict(self):
        return {
            "left": self.left,
            "top": self.top,
            "width": self.width,
            "height": self.height
        }

    @classmethod
    def from_dict(cls, data: dict):
        return cls(
            left=data["left"],
            top=data["top"],
            width=data["width"],
            height=data["height"]
        )


@dataclass
class NormalizedGeometry:
    """Geometry information with normalized coordinates"""
    bounding_box: BoundingBox
    polygon: list  # List[Point] - corners of the text region

    def to_dict(self):
        return {
            "bounding_box": self.bounding_box.to_dict(),
            "polygon": [p.to_dict() for p in self.polygon]
        }

    @classmethod
    def from_dict(cls, data: dict):
        return cls(
            bounding_box=BoundingBox.from_dict(data["bounding_box"]),
            polygon=[Point.from_dict(p) for p in data["polygon"]]
        )


@dataclass
class BlockRelationship:
    """Relationship between blocks (e.g., LINE contains WORDs)"""
    type: str  # "CHILD", "VALUE", "TITLE"
    ids: list  # List of related block IDs

    def to_dict(self):
        return {"type": self.type, "ids": self.ids}

    @classmethod
    def from_dict(cls, data: dict):
        return cls(type=data["type"], ids=data["ids"])


@dataclass
class NormalizedBlock:
    """
    Internal block format (inspired by AWS Textract structure).
    Represents a single text element (word, line, page, table, cell).
    """
    id: str
    block_type: str  # WORD, LINE, PAGE, TABLE, CELL
    text: Optional[str]
    confidence: float  # 0-100
    geometry: NormalizedGeometry
    page: int
    relationships: list = None  # List[BlockRelationship]

    def __post_init__(self):
        if self.relationships is None:
            self.relationships = []

    def to_dict(self):
        return {
            "id": self.id,
            "block_type": self.block_type,
            "text": self.text,
            "confidence": self.confidence,
            "geometry": self.geometry.to_dict(),
            "page": self.page,
            "relationships": [r.to_dict() for r in self.relationships]
        }

    @classmethod
    def from_dict(cls, data: dict):
        return cls(
            id=data["id"],
            block_type=data["block_type"],
            text=data.get("text"),
            confidence=data["confidence"],
            geometry=NormalizedGeometry.from_dict(data["geometry"]),
            page=data["page"],
            relationships=[BlockRelationship.from_dict(r) for r in data.get("relationships", [])]
        )


@dataclass
class DocumentMetadata:
    """Metadata about the processed document"""
    pages: int
    extraction_method: Optional[str] = None
    processing_time_ms: Optional[int] = None

    def to_dict(self):
        return {
            "pages": self.pages,
            "extraction_method": self.extraction_method,
            "processing_time_ms": self.processing_time_ms
        }

    @classmethod
    def from_dict(cls, data: dict):
        return cls(
            pages=data["pages"],
            extraction_method=data.get("extraction_method"),
            processing_time_ms=data.get("processing_time_ms")
        )


@dataclass
class NormalizedResult:
    """
    Complete extraction result in normalized format.
    This is the internal format used by DTAT - all formatters convert from this.
    """
    blocks: list  # List[NormalizedBlock]
    document_metadata: DocumentMetadata
    page_count: int
    confidence_score: float  # Overall 0-100

    def to_dict(self):
        return {
            "blocks": [b.to_dict() for b in self.blocks],
            "document_metadata": self.document_metadata.to_dict(),
            "page_count": self.page_count,
            "confidence_score": self.confidence_score
        }

    @classmethod
    def from_dict(cls, data: dict):
        return cls(
            blocks=[NormalizedBlock.from_dict(b) for b in data["blocks"]],
            document_metadata=DocumentMetadata.from_dict(data["document_metadata"]),
            page_count=data["page_count"],
            confidence_score=data["confidence_score"]
        )


# ==============================================================================
# Conversion Helper
# ==============================================================================

def convert_extraction_result_to_normalized(
    extraction_result,
    page_count: int = 1
) -> NormalizedResult:
    """
    Convert legacy ExtractionResult to new NormalizedResult format.

    This maintains backward compatibility while enabling multi-format output.

    Args:
        extraction_result: Legacy ExtractionResult object
        page_count: Number of pages in document

    Returns:
        NormalizedResult with blocks created from text lines
    """
    blocks = []
    block_id = 0

    # Split text into lines and create LINE blocks
    text_lines = extraction_result.text_content.split('\n')

    for line_idx, line_text in enumerate(text_lines):
        if not line_text.strip():
            continue  # Skip empty lines

        # Create a LINE block
        # Note: We don't have actual coordinates yet, so we approximate
        # based on line position. Later we can enhance extraction to provide real coords.
        line_height = 1.0 / max(len(text_lines), 1)
        top_position = line_idx * line_height

        # Create simple bounding box (full width, proportional height)
        bbox = BoundingBox(
            left=0.05,  # 5% margin from left
            top=top_position,
            width=0.90,  # 90% width (5% margins on each side)
            height=line_height
        )

        # Create polygon (4 corners of bounding box)
        polygon = [
            Point(bbox.left, bbox.top),
            Point(bbox.left + bbox.width, bbox.top),
            Point(bbox.left + bbox.width, bbox.top + bbox.height),
            Point(bbox.left, bbox.top + bbox.height)
        ]

        geometry = NormalizedGeometry(
            bounding_box=bbox,
            polygon=polygon
        )

        # Determine which page this line is on (rough approximation)
        lines_per_page = max(len(text_lines) / page_count, 1) if page_count > 1 else len(text_lines)
        current_page = min(int(line_idx / lines_per_page) + 1, page_count)

        block = NormalizedBlock(
            id=f"block_{block_id}",
            block_type="LINE",
            text=line_text,
            confidence=extraction_result.confidence_score,
            geometry=geometry,
            page=current_page,
            relationships=[]
        )

        blocks.append(block)
        block_id += 1

    # Create PAGE blocks (one per page)
    for page_num in range(1, page_count + 1):
        page_block = NormalizedBlock(
            id=f"page_{page_num}",
            block_type="PAGE",
            text=None,
            confidence=extraction_result.confidence_score,
            geometry=NormalizedGeometry(
                bounding_box=BoundingBox(0, 0, 1, 1),  # Full page
                polygon=[Point(0, 0), Point(1, 0), Point(1, 1), Point(0, 1)]
            ),
            page=page_num,
            relationships=[]
        )
        blocks.append(page_block)

    # Create document metadata
    metadata = DocumentMetadata(
        pages=page_count,
        extraction_method=extraction_result.method_used,
        processing_time_ms=extraction_result.processing_time_ms
    )

    # Create normalized result
    return NormalizedResult(
        blocks=blocks,
        document_metadata=metadata,
        page_count=page_count,
        confidence_score=extraction_result.confidence_score
    )


# ==============================================================================
# Legacy Extraction Result (for backward compatibility)
# ==============================================================================

@dataclass
class ExtractionResult:
    """Result from any extraction method."""
    success: bool
    text_content: str
    tables: list
    metadata: dict
    confidence_score: float
    method_used: str
    error_message: Optional[str] = None
    processing_time_ms: int = 0


class QualityScorer:
    """Evaluate extraction quality to decide if we need to retry."""

    @staticmethod
    def calculate_gibberish_ratio(text: str) -> float:
        """Calculate ratio of non-readable characters."""
        if not text:
            return 1.0

        # Count "normal" characters (letters, numbers, common punctuation, spaces)
        normal_pattern = re.compile(r'[a-zA-Z0-9\s\.,;:!?\'"()\-\n\t@#$%&*+=/<>]')
        normal_chars = len(normal_pattern.findall(text))
        total_chars = len(text)

        if total_chars == 0:
            return 1.0

        return 1.0 - (normal_chars / total_chars)

    @staticmethod
    def has_expected_structure(text: str) -> bool:
        """Check if text has sentence-like structure."""
        # Look for patterns like sentences, paragraphs
        sentences = re.findall(r'[A-Z][^.!?]*[.!?]', text)
        return len(sentences) >= 2

    @staticmethod
    def calculate_confidence(result: ExtractionResult, page_count: int = 1) -> float:
        """
        Calculate confidence score (0-100) for extraction quality.
        Higher score = better quality extraction.
        """
        score = 0.0
        text = result.text_content

        if not text:
            return 0.0

        # 1. Text length score (0-30 points)
        chars_per_page = len(text) / max(page_count, 1)
        if chars_per_page >= 500:
            score += 30
        elif chars_per_page >= 200:
            score += 20
        elif chars_per_page >= 100:
            score += 10
        elif chars_per_page >= 50:
            score += 5

        # 2. Low gibberish score (0-25 points)
        gibberish_ratio = QualityScorer.calculate_gibberish_ratio(text)
        if gibberish_ratio < 0.05:
            score += 25
        elif gibberish_ratio < 0.10:
            score += 20
        elif gibberish_ratio < 0.15:
            score += 15
        elif gibberish_ratio < 0.25:
            score += 10

        # 3. Has structure (0-20 points)
        if QualityScorer.has_expected_structure(text):
            score += 20
        elif len(text.split('\n')) > 3:  # At least has line breaks
            score += 10

        # 4. Tables found bonus (0-15 points)
        if result.tables:
            score += 15

        # 5. No error (0-10 points)
        if not result.error_message:
            score += 10

        return min(score, 100.0)


# =============================================================================
# EXTRACTION METHODS
# =============================================================================

class NativeExtractor:
    """Level 1: Free extraction using standard libraries."""

    @staticmethod
    def can_handle(file_type: str) -> bool:
        return file_type in ['pdf', 'xlsx', 'xls', 'csv', 'docx', 'doc']

    @staticmethod
    def extract(file_path: Path, file_type: str) -> ExtractionResult:
        start_time = time.time()

        try:
            if file_type == 'pdf':
                return NativeExtractor._extract_pdf(file_path, start_time)
            elif file_type in ['xlsx', 'xls']:
                return NativeExtractor._extract_excel(file_path, start_time)
            elif file_type == 'csv':
                return NativeExtractor._extract_csv(file_path, start_time)
            elif file_type in ['docx', 'doc']:
                return NativeExtractor._extract_word(file_path, start_time)
            else:
                return ExtractionResult(
                    success=False,
                    text_content="",
                    tables=[],
                    metadata={},
                    confidence_score=0,
                    method_used=ExtractionMethod.NATIVE.value,
                    error_message=f"Unsupported file type for native extraction: {file_type}"
                )
        except Exception as e:
            return ExtractionResult(
                success=False,
                text_content="",
                tables=[],
                metadata={},
                confidence_score=0,
                method_used=ExtractionMethod.NATIVE.value,
                error_message=str(e),
                processing_time_ms=int((time.time() - start_time) * 1000)
            )

    @staticmethod
    def _extract_pdf(file_path: Path, start_time: float) -> ExtractionResult:
        import pdfplumber

        # Check for password-protected/encrypted PDFs before processing
        try:
            with pdfplumber.open(file_path) as test_pdf:
                # Must actually access a page — some encrypted PDFs open but fail on page access
                if len(test_pdf.pages) > 0:
                    test_pdf.pages[0].extract_text()
        except Exception as e:
            err_str = str(e).lower()
            if any(kw in err_str for kw in ['encrypt', 'password', 'protected', 'permission', 'crypt']):
                return ExtractionResult(
                    success=False, text_content="", tables=[], metadata={"is_encrypted": True},
                    confidence_score=0, method_used=ExtractionMethod.NATIVE.value,
                    error_message="PDF is password-protected. Please provide an unencrypted version.",
                    processing_time_ms=int((time.time() - start_time) * 1000)
                )
            raise

        all_text = []
        all_tables = []

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)

            for page in pdf.pages:
                # Extract text
                text = page.extract_text() or ""
                all_text.append(text)

                # Extract tables
                tables = page.extract_tables()
                for table in tables:
                    if table and len(table) > 1:
                        # First row as headers, rest as data
                        headers = table[0] if table[0] else [f"col_{i}" for i in range(len(table[1]))]
                        records = []
                        for row in table[1:]:
                            if row:
                                records.append(dict(zip(headers, row)))
                        if records:
                            all_tables.append(records)

        full_text = "\n\n".join(all_text)

        result = ExtractionResult(
            success=True,
            text_content=full_text,
            tables=all_tables,
            metadata={"pages": page_count},
            confidence_score=0,  # Will be calculated later
            method_used=ExtractionMethod.NATIVE.value,
            processing_time_ms=int((time.time() - start_time) * 1000)
        )

        result.confidence_score = QualityScorer.calculate_confidence(result, page_count)
        return result

    @staticmethod
    def _extract_excel(file_path: Path, start_time: float) -> ExtractionResult:
        import pandas as pd

        excel_file = pd.ExcelFile(file_path)
        all_text = []
        all_tables = []

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            all_text.append(f"Sheet: {sheet_name}\n{df.to_string()}")
            all_tables.append({
                "sheet": sheet_name,
                "data": df.to_dict(orient='records')
            })

        return ExtractionResult(
            success=True,
            text_content="\n\n".join(all_text),
            tables=all_tables,
            metadata={"sheets": excel_file.sheet_names},
            confidence_score=95,  # Excel extraction is reliable
            method_used=ExtractionMethod.NATIVE.value,
            processing_time_ms=int((time.time() - start_time) * 1000)
        )

    @staticmethod
    def _extract_csv(file_path: Path, start_time: float) -> ExtractionResult:
        import pandas as pd

        df = pd.read_csv(file_path)

        return ExtractionResult(
            success=True,
            text_content=df.to_string(),
            tables=[{"data": df.to_dict(orient='records')}],
            metadata={"columns": list(df.columns), "rows": len(df)},
            confidence_score=98,  # CSV extraction is very reliable
            method_used=ExtractionMethod.NATIVE.value,
            processing_time_ms=int((time.time() - start_time) * 1000)
        )

    @staticmethod
    def _extract_word(file_path: Path, start_time: float) -> ExtractionResult:
        from docx import Document

        doc = Document(file_path)
        all_text = [para.text for para in doc.paragraphs if para.text.strip()]
        all_tables = []

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            if table_data and len(table_data) > 1:
                headers = table_data[0]
                records = [dict(zip(headers, row)) for row in table_data[1:]]
                all_tables.append(records)

        return ExtractionResult(
            success=True,
            text_content="\n\n".join(all_text),
            tables=all_tables,
            metadata={},
            confidence_score=95,  # Word extraction is reliable
            method_used=ExtractionMethod.NATIVE.value,
            processing_time_ms=int((time.time() - start_time) * 1000)
        )


class LocalOCRExtractor:
    """Level 2: LightOnOCR for scanned documents."""

    _model = None
    _processor = None
    _device = None
    _dtype = None

    @classmethod
    def _load_model(cls):
        """Lazy load the OCR model."""
        if cls._model is None:
            import torch
            from transformers import LightOnOcrForConditionalGeneration, LightOnOcrProcessor

            if torch.cuda.is_available():
                cls._device = "cuda"
                cls._dtype = torch.bfloat16
            elif torch.backends.mps.is_available():
                cls._device = "mps"
                cls._dtype = torch.float32
            else:
                cls._device = "cpu"
                cls._dtype = torch.float32

            print(f"Loading LightOnOCR model (device: {cls._device}, offline: {config.ocr_offline_mode})...")
            cls._model = LightOnOcrForConditionalGeneration.from_pretrained(
                config.ocr_model_name,
                torch_dtype=cls._dtype,
                local_files_only=config.ocr_offline_mode
            ).to(cls._device)
            cls._processor = LightOnOcrProcessor.from_pretrained(
                config.ocr_model_name,
                local_files_only=config.ocr_offline_mode
            )
            print("Model loaded.")

    @classmethod
    def extract(cls, file_path: Path, file_type: str) -> ExtractionResult:
        from PIL import Image
        import pypdfium2 as pdfium

        start_time = time.time()

        try:
            cls._load_model()

            all_text = []
            page_count = 1

            if file_type == 'pdf':
                pdf = pdfium.PdfDocument(file_path)
                page_count = len(pdf)

                for i in range(page_count):
                    print(f"  OCR page {i + 1}/{page_count}...")
                    page = pdf[i]
                    pil_image = page.render(scale=2.77).to_pil()  # 200 DPI
                    text = cls._ocr_single_image(pil_image)
                    all_text.append(text)
            else:
                # Image file
                image = Image.open(file_path)
                text = cls._ocr_single_image(image)
                all_text.append(text)

            full_text = "\n\n".join(all_text)

            result = ExtractionResult(
                success=True,
                text_content=full_text,
                tables=[],  # OCR doesn't extract structured tables
                metadata={"pages": page_count, "ocr_device": cls._device},
                confidence_score=0,
                method_used=ExtractionMethod.LOCAL_OCR.value,
                processing_time_ms=int((time.time() - start_time) * 1000)
            )

            result.confidence_score = QualityScorer.calculate_confidence(result, page_count)
            return result

        except Exception as e:
            return ExtractionResult(
                success=False,
                text_content="",
                tables=[],
                metadata={},
                confidence_score=0,
                method_used=ExtractionMethod.LOCAL_OCR.value,
                error_message=f"{type(e).__name__}: {str(e)}",
                processing_time_ms=int((time.time() - start_time) * 1000)
            )

    @classmethod
    def _ocr_single_image(cls, image) -> str:
        from PIL import Image

        # Resize to recommended max dimension
        max_dim = config.ocr_image_max_dim
        if max(image.size) > max_dim:
            ratio = max_dim / max(image.size)
            new_size = (int(image.size[0] * ratio), int(image.size[1] * ratio))
            image = image.resize(new_size, Image.Resampling.LANCZOS)

        conversation = [{"role": "user", "content": [{"type": "image", "image": image}]}]
        inputs = cls._processor.apply_chat_template(
            conversation,
            add_generation_prompt=True,
            tokenize=True,
            return_dict=True,
            return_tensors="pt",
        )
        inputs = {
            k: v.to(device=cls._device, dtype=cls._dtype) if v.is_floating_point() else v.to(cls._device)
            for k, v in inputs.items()
        }

        output_ids = cls._model.generate(**inputs, max_new_tokens=config.ocr_max_new_tokens)
        generated_ids = output_ids[0, inputs["input_ids"].shape[1]:]
        return cls._processor.decode(generated_ids, skip_special_tokens=True)


class TextractExtractor:
    """Level 3: AWS Textract (paid, optional)."""

    _client = None

    @classmethod
    def _get_client(cls):
        """Reuse boto3 Textract client across requests with adaptive retry."""
        if cls._client is None:
            import boto3
            from botocore.config import Config
            retry_config = Config(
                retries={'max_attempts': 5, 'mode': 'adaptive'},
                read_timeout=60,
                connect_timeout=10,
            )
            cls._client = boto3.client('textract', region_name=config.aws_region, config=retry_config)
        return cls._client

    _s3_client = None

    @classmethod
    def _get_s3_client(cls):
        """Reuse boto3 S3 client."""
        if cls._s3_client is None:
            import boto3
            cls._s3_client = boto3.client('s3', region_name=config.aws_region)
        return cls._s3_client

    @classmethod
    def _process_via_s3(cls, file_path: Path, file_bytes: bytes) -> list:
        """Upload to S3, run async Textract, poll for results, cleanup."""
        import uuid as _uuid

        s3 = cls._get_s3_client()
        client = cls._get_client()

        s3_key = f"{config.s3_prefix}{_uuid.uuid4()}.pdf"

        try:
            # Upload to S3
            s3.put_object(Bucket=config.s3_bucket, Key=s3_key, Body=file_bytes)

            # Start async text detection
            response = client.start_document_text_detection(
                DocumentLocation={'S3Object': {'Bucket': config.s3_bucket, 'Name': s3_key}}
            )
            job_id = response['JobId']

            # Poll until complete (max 5 minutes, 5s intervals)
            for _ in range(60):
                result = client.get_document_text_detection(JobId=job_id)
                status = result['JobStatus']
                if status == 'SUCCEEDED':
                    all_blocks = result.get('Blocks', [])
                    # Handle pagination for large documents
                    next_token = result.get('NextToken')
                    while next_token:
                        result = client.get_document_text_detection(JobId=job_id, NextToken=next_token)
                        all_blocks.extend(result.get('Blocks', []))
                        next_token = result.get('NextToken')
                    return all_blocks
                elif status == 'FAILED':
                    raise Exception(f"Textract async job failed: {result.get('StatusMessage', 'Unknown error')}")
                time.sleep(5)

            raise Exception("Textract async job timed out after 5 minutes")

        finally:
            # Cleanup S3 temp file
            try:
                s3.delete_object(Bucket=config.s3_bucket, Key=s3_key)
            except Exception:
                pass

    @classmethod
    def extract(cls, file_path: Path, file_type: str) -> ExtractionResult:
        start_time = time.time()

        if not config.enable_textract:
            return ExtractionResult(
                success=False,
                text_content="",
                tables=[],
                metadata={},
                confidence_score=0,
                method_used=ExtractionMethod.TEXTRACT.value,
                error_message="Textract is disabled in config",
                processing_time_ms=int((time.time() - start_time) * 1000)
            )

        try:
            client = cls._get_client()

            with open(file_path, 'rb') as f:
                file_bytes = f.read()

            all_blocks = []

            if file_type == 'pdf' and len(file_bytes) > 5 * 1024 * 1024:
                # Large PDFs (>5MB): use S3 upload + async Textract API
                if not config.s3_bucket:
                    return ExtractionResult(
                        success=False, text_content="", tables=[], metadata={},
                        confidence_score=0, method_used=ExtractionMethod.TEXTRACT.value,
                        error_message="PDF too large for direct Textract (>5MB). Set S3_BUCKET env var to enable large file processing.",
                        processing_time_ms=int((time.time() - start_time) * 1000)
                    )
                all_blocks = cls._process_via_s3(file_path, file_bytes)
            elif file_type == 'pdf':
                # PDFs: use analyze_document with TABLES+FORMS for richer extraction
                # Textract sync API handles multi-page PDFs up to 5MB
                response = client.analyze_document(
                    Document={'Bytes': file_bytes},
                    FeatureTypes=config.textract_features
                )
                all_blocks = response.get('Blocks', [])
            else:
                # Images: use detect_document_text (simpler, faster)
                response = client.detect_document_text(
                    Document={'Bytes': file_bytes}
                )
                all_blocks = response.get('Blocks', [])

            # Extract text from LINE blocks across all pages
            all_text = []
            page_count = 0
            for block in all_blocks:
                if block['BlockType'] == 'LINE':
                    all_text.append(block.get('Text', ''))
                if block['BlockType'] == 'PAGE':
                    page_count += 1

            # Extract tables if present
            tables = []
            table_blocks = [b for b in all_blocks if b['BlockType'] == 'TABLE']
            if table_blocks:
                tables = [{"table_id": b.get('Id', ''), "cells": len(b.get('Relationships', []))} for b in table_blocks]

            result = ExtractionResult(
                success=True,
                text_content="\n".join(all_text),
                tables=tables,
                metadata={
                    "textract_blocks": len(all_blocks),
                    "pages": page_count or 1,
                    "tables_found": len(tables),
                },
                confidence_score=90,
                method_used=ExtractionMethod.TEXTRACT.value,
                processing_time_ms=int((time.time() - start_time) * 1000)
            )

            return result

        except Exception as e:
            err_msg = str(e).lower()
            # Detect encrypted/unsupported PDFs at Textract level
            if 'unsupporteddocument' in err_msg or 'encrypt' in err_msg:
                if file_type == 'pdf':
                    return ExtractionResult(
                        success=False, text_content="", tables=[],
                        metadata={"is_encrypted": True},
                        confidence_score=0, method_used=ExtractionMethod.TEXTRACT.value,
                        error_message="PDF is password-protected or has unsupported encryption. Please provide an unencrypted version.",
                        processing_time_ms=int((time.time() - start_time) * 1000)
                    )
            return ExtractionResult(
                success=False,
                text_content="",
                tables=[],
                metadata={},
                confidence_score=0,
                method_used=ExtractionMethod.TEXTRACT.value,
                error_message=f"{type(e).__name__}: {str(e)}",
                processing_time_ms=int((time.time() - start_time) * 1000)
            )


# =============================================================================
# MAIN PIPELINE
# =============================================================================

class ExtractionPipeline:
    """
    Main extraction pipeline with retry logic and escalation.
    """

    def __init__(self, cfg: ProcessingConfig = None):
        self.config = cfg or config

    def process(self, document: DocumentRecord, file_path: Path) -> DocumentRecord:
        """
        Process a document through the extraction ladder.
        Returns updated document record.
        """
        document.status = ProcessingStatus.PROCESSING.value
        document.started_at = datetime.utcnow()
        update_document(document)

        file_type = document.file_type.lower().lstrip('.')
        is_image = file_type in ['jpg', 'jpeg', 'png', 'tiff', 'tif', 'bmp', 'gif', 'webp']
        levels_tried = []
        attempt_number = 0

        # Level 1: Native extraction (skip for images)
        if self.config.enable_native_extraction and not is_image:
            for retry in range(self.config.max_retries_per_level):
                attempt_number += 1
                result = self._try_extraction(
                    document, file_path, file_type,
                    NativeExtractor.extract, "native", attempt_number
                )
                levels_tried.append(f"native_attempt_{retry + 1}")

                if result.success and result.confidence_score >= self.config.min_confidence_score:
                    return self._finalize_success(document, result, levels_tried)

                if retry < self.config.max_retries_per_level - 1:
                    time.sleep(self.config.retry_delay_seconds)

        # Level 2: Local OCR
        if self.config.enable_local_ocr:
            for retry in range(self.config.max_retries_per_level):
                attempt_number += 1
                result = self._try_extraction(
                    document, file_path, file_type,
                    LocalOCRExtractor.extract, "local_ocr", attempt_number
                )
                levels_tried.append(f"local_ocr_attempt_{retry + 1}")

                if result.success and result.confidence_score >= self.config.min_confidence_score:
                    return self._finalize_success(document, result, levels_tried)

                if retry < self.config.max_retries_per_level - 1:
                    time.sleep(self.config.retry_delay_seconds)

        # Level 3: Textract (if enabled)
        if self.config.enable_textract:
            for retry in range(self.config.max_retries_per_level):
                attempt_number += 1
                result = self._try_extraction(
                    document, file_path, file_type,
                    TextractExtractor.extract, "textract", attempt_number
                )
                levels_tried.append(f"textract_attempt_{retry + 1}")

                if result.success and result.confidence_score >= self.config.min_confidence_score:
                    return self._finalize_success(document, result, levels_tried)

                if retry < self.config.max_retries_per_level - 1:
                    time.sleep(self.config.retry_delay_seconds)

        # All levels exhausted - send to DLQ
        return self._finalize_failure(document, result, levels_tried)

    def _try_extraction(
        self,
        document: DocumentRecord,
        file_path: Path,
        file_type: str,
        extractor_fn,
        method_name: str,
        attempt_number: int
    ) -> ExtractionResult:
        """Try a single extraction method and log the attempt."""
        try:
            result = extractor_fn(file_path, file_type)
        except Exception as e:
            result = ExtractionResult(
                success=False,
                text_content="",
                tables=[],
                metadata={},
                confidence_score=0,
                method_used=method_name,
                error_message=f"Unhandled exception: {traceback.format_exc()}"
            )

        # Log the attempt
        log_processing_attempt(
            document_id=document.id,
            attempt_number=attempt_number,
            method=method_name,
            success=result.success and result.confidence_score >= self.config.min_confidence_score,
            duration_ms=result.processing_time_ms,
            confidence=result.confidence_score,
            chars=len(result.text_content),
            tables=len(result.tables),
            error=result.error_message
        )

        return result

    def _finalize_success(
        self,
        document: DocumentRecord,
        result: ExtractionResult,
        levels_tried: list
    ) -> DocumentRecord:
        """Finalize a successful extraction."""
        document.status = ProcessingStatus.COMPLETED.value
        document.completed_at = datetime.utcnow()
        document.extraction_method = result.method_used
        document.extraction_levels_tried = str(levels_tried)
        document.confidence_score = result.confidence_score
        document.processing_time_ms = result.processing_time_ms
        document.char_count = len(result.text_content)
        document.table_count = len(result.tables)
        document.page_count = result.metadata.get('pages', 1)

        # Convert to normalized format and store
        normalized_result = convert_extraction_result_to_normalized(
            result,
            page_count=document.page_count
        )
        document.set_normalized_content(normalized_result)

        # Profile-based extraction (if profile assigned)
        if document.profile_id:
            self._extract_with_profile(document, normalized_result)

        update_document(document)
        return document

    def _finalize_failure(
        self,
        document: DocumentRecord,
        last_result: ExtractionResult,
        levels_tried: list
    ) -> DocumentRecord:
        """Finalize a failed extraction - send to DLQ."""
        document.status = ProcessingStatus.NEEDS_REVIEW.value
        document.completed_at = datetime.utcnow()
        document.extraction_levels_tried = str(levels_tried)
        document.error_message = f"All extraction methods failed. Last error: {last_result.error_message}"
        document.retry_count = len(levels_tried)

        update_document(document)
        print(f"[DLQ] Document {document.id} needs manual review: {document.error_message}")
        return document

    def _extract_with_profile(
        self,
        document: DocumentRecord,
        ocr_result: dict
    ) -> None:
        """
        Extract structured fields using assigned profile.

        Args:
            document: Document record with profile_id set
            ocr_result: Normalized OCR result from extraction

        Side effects:
            - Sets document.extracted_fields (JSONB)
            - Logs to profile_usage table
        """
        from database import get_profile_by_id, log_profile_usage
        from extractors import ProfileExtractor
        from profiles import ExtractionProfile

        start_time = time.time()

        try:
            # Get profile
            profile_record = get_profile_by_id(document.profile_id)
            if not profile_record:
                print(f"[WARN] Profile {document.profile_id} not found for document {document.id}")
                return

            # Convert record to ExtractionProfile
            schema = profile_record.get_schema()
            schema['id'] = profile_record.id
            profile = ExtractionProfile(**schema)

            # Extract fields
            extractor = ProfileExtractor()
            extraction_results = extractor.extract_all_fields(profile, ocr_result)

            # Store extracted fields
            document.extracted_fields = extraction_results

            # Calculate statistics
            stats = extraction_results['statistics']
            processing_time_ms = int((time.time() - start_time) * 1000)

            # Determine status
            if stats['failed'] == 0 and stats['extracted'] >= stats['required']:
                status = 'success'
            elif stats['extracted'] > 0:
                status = 'partial'
            else:
                status = 'failed'

            # Log usage
            log_profile_usage(
                profile_id=document.profile_id,
                document_id=document.id,
                fields_extracted=stats['extracted'],
                fields_failed=stats['failed'],
                avg_confidence=stats.get('avg_confidence', 0.0),
                processing_time_ms=processing_time_ms,
                status=status
            )

            print(f"[PROFILE] Extracted {stats['extracted']}/{stats['total_fields']} fields from document {document.id}")

        except Exception as e:
            error_msg = f"Profile extraction failed: {str(e)}"
            print(f"[ERROR] {error_msg}")
            print(traceback.format_exc())

            # Log failed usage
            try:
                processing_time_ms = int((time.time() - start_time) * 1000)
                log_profile_usage(
                    profile_id=document.profile_id,
                    document_id=document.id,
                    fields_extracted=0,
                    fields_failed=0,
                    avg_confidence=0.0,
                    processing_time_ms=processing_time_ms,
                    status='failed',
                    error_message=error_msg
                )
            except:
                pass  # Don't fail if logging fails
