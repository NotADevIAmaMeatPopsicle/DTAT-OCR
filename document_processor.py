"""
Swiss Army Knife Document Processor
Handles: Scanned images, PDFs (native + scanned), Excel, CSV, Word

Routes to the right tool based on document type:
- Scanned/Image → LightOnOCR (AI-based OCR)
- Native PDF → PyMuPDF (direct text extraction)
- Excel → openpyxl/pandas
- CSV → pandas
- Word → python-docx
"""

import sys
import json
from pathlib import Path
from typing import Any
from dataclasses import dataclass, asdict
from enum import Enum

# Document processing libraries
import fitz  # PyMuPDF
import pandas as pd
from openpyxl import load_workbook
from docx import Document as DocxDocument


class DocumentType(Enum):
    IMAGE = "image"
    PDF_NATIVE = "pdf_native"
    PDF_SCANNED = "pdf_scanned"
    EXCEL = "excel"
    CSV = "csv"
    WORD = "word"
    UNKNOWN = "unknown"


@dataclass
class ProcessedDocument:
    """Structured output for any processed document."""
    source_file: str
    document_type: str
    text_content: str
    tables: list[dict]  # List of tables as list of dicts
    metadata: dict
    pages: int
    used_ocr: bool


def detect_document_type(file_path: Path) -> DocumentType:
    """Detect document type from file extension and content."""
    suffix = file_path.suffix.lower()

    if suffix in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif', '.webp']:
        return DocumentType.IMAGE
    elif suffix == '.pdf':
        # Check if PDF has extractable text or is scanned
        return _check_pdf_type(file_path)
    elif suffix in ['.xlsx', '.xls']:
        return DocumentType.EXCEL
    elif suffix == '.csv':
        return DocumentType.CSV
    elif suffix in ['.docx', '.doc']:
        return DocumentType.WORD
    else:
        return DocumentType.UNKNOWN


def _check_pdf_type(file_path: Path) -> DocumentType:
    """Check if PDF is native (has text) or scanned (needs OCR)."""
    try:
        doc = fitz.open(file_path)
        total_text = ""
        for page in doc:
            total_text += page.get_text()
        doc.close()

        # If we got meaningful text, it's a native PDF
        # Threshold: at least 50 chars per page on average
        avg_chars = len(total_text.strip()) / max(len(doc), 1)
        if avg_chars > 50:
            return DocumentType.PDF_NATIVE
        else:
            return DocumentType.PDF_SCANNED
    except Exception:
        return DocumentType.PDF_SCANNED


# =============================================================================
# NATIVE DOCUMENT PROCESSORS (No OCR needed)
# =============================================================================

def process_native_pdf(file_path: Path) -> ProcessedDocument:
    """Extract text from a native PDF using PyMuPDF."""
    doc = fitz.open(file_path)

    all_text = []
    all_tables = []

    for page_num, page in enumerate(doc):
        # Extract text
        text = page.get_text()
        all_text.append(f"--- Page {page_num + 1} ---\n{text}")

        # Extract tables (PyMuPDF can find tables)
        tables = page.find_tables()
        for table in tables:
            df = table.to_pandas()
            all_tables.append(df.to_dict(orient='records'))

    metadata = {
        "title": doc.metadata.get("title", ""),
        "author": doc.metadata.get("author", ""),
        "creator": doc.metadata.get("creator", ""),
        "producer": doc.metadata.get("producer", ""),
    }

    page_count = len(doc)
    doc.close()

    return ProcessedDocument(
        source_file=str(file_path),
        document_type="pdf_native",
        text_content="\n\n".join(all_text),
        tables=all_tables,
        metadata=metadata,
        pages=page_count,
        used_ocr=False
    )


def process_excel(file_path: Path) -> ProcessedDocument:
    """Extract data from Excel files."""
    # Read all sheets
    excel_file = pd.ExcelFile(file_path)

    all_text = []
    all_tables = []

    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(excel_file, sheet_name=sheet_name)

        all_text.append(f"--- Sheet: {sheet_name} ---")
        all_text.append(df.to_string())

        # Convert to records for structured output
        all_tables.append({
            "sheet": sheet_name,
            "data": df.to_dict(orient='records')
        })

    return ProcessedDocument(
        source_file=str(file_path),
        document_type="excel",
        text_content="\n\n".join(all_text),
        tables=all_tables,
        metadata={"sheets": excel_file.sheet_names},
        pages=len(excel_file.sheet_names),
        used_ocr=False
    )


def process_csv(file_path: Path) -> ProcessedDocument:
    """Extract data from CSV files."""
    df = pd.read_csv(file_path)

    return ProcessedDocument(
        source_file=str(file_path),
        document_type="csv",
        text_content=df.to_string(),
        tables=[{"data": df.to_dict(orient='records')}],
        metadata={"columns": list(df.columns), "rows": len(df)},
        pages=1,
        used_ocr=False
    )


def process_word(file_path: Path) -> ProcessedDocument:
    """Extract text and tables from Word documents."""
    doc = DocxDocument(file_path)

    all_text = []
    all_tables = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text)

    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)

        if table_data:
            # Convert to dict format (first row as headers)
            if len(table_data) > 1:
                headers = table_data[0]
                records = [dict(zip(headers, row)) for row in table_data[1:]]
                all_tables.append(records)
            else:
                all_tables.append(table_data)

    # Get metadata
    core_props = doc.core_properties
    metadata = {
        "title": core_props.title or "",
        "author": core_props.author or "",
        "created": str(core_props.created) if core_props.created else "",
        "modified": str(core_props.modified) if core_props.modified else "",
    }

    return ProcessedDocument(
        source_file=str(file_path),
        document_type="word",
        text_content="\n\n".join(all_text),
        tables=all_tables,
        metadata=metadata,
        pages=1,  # Word doesn't have fixed pages
        used_ocr=False
    )


# =============================================================================
# OCR PROCESSORS (For scanned documents)
# =============================================================================

def process_with_ocr(file_path: Path, doc_type: DocumentType) -> ProcessedDocument:
    """Process scanned documents using LightOnOCR."""
    # Lazy import to avoid loading the model unless needed
    import torch
    from transformers import LightOnOcrForConditionalGeneration, LightOnOcrProcessor
    from PIL import Image
    import pypdfium2 as pdfium

    def get_device_and_dtype():
        if torch.cuda.is_available():
            return "cuda", torch.bfloat16
        elif torch.backends.mps.is_available():
            return "mps", torch.float32
        else:
            return "cpu", torch.float32

    device, dtype = get_device_and_dtype()
    print(f"Loading OCR model (device: {device})...")

    model = LightOnOcrForConditionalGeneration.from_pretrained(
        "lightonai/LightOnOCR-1B-1025",
        torch_dtype=dtype
    ).to(device)
    processor = LightOnOcrProcessor.from_pretrained("lightonai/LightOnOCR-1B-1025")

    def ocr_image(image: Image.Image) -> str:
        # Resize to recommended max dimension
        max_dim = 1540
        if max(image.size) > max_dim:
            ratio = max_dim / max(image.size)
            new_size = (int(image.size[0] * ratio), int(image.size[1] * ratio))
            image = image.resize(new_size, Image.Resampling.LANCZOS)

        conversation = [{"role": "user", "content": [{"type": "image", "image": image}]}]
        inputs = processor.apply_chat_template(
            conversation,
            add_generation_prompt=True,
            tokenize=True,
            return_dict=True,
            return_tensors="pt",
        )
        inputs = {
            k: v.to(device=device, dtype=dtype) if v.is_floating_point() else v.to(device)
            for k, v in inputs.items()
        }

        output_ids = model.generate(**inputs, max_new_tokens=2048)
        generated_ids = output_ids[0, inputs["input_ids"].shape[1]:]
        return processor.decode(generated_ids, skip_special_tokens=True)

    all_text = []
    pages = 1

    if doc_type == DocumentType.IMAGE:
        image = Image.open(file_path)
        text = ocr_image(image)
        all_text.append(text)

    elif doc_type == DocumentType.PDF_SCANNED:
        pdf = pdfium.PdfDocument(file_path)
        pages = len(pdf)

        for i in range(pages):
            print(f"  OCR processing page {i + 1}/{pages}...")
            page = pdf[i]
            pil_image = page.render(scale=2.77).to_pil()  # 200 DPI
            text = ocr_image(pil_image)
            all_text.append(f"--- Page {i + 1} ---\n{text}")

    return ProcessedDocument(
        source_file=str(file_path),
        document_type=doc_type.value,
        text_content="\n\n".join(all_text),
        tables=[],  # OCR doesn't extract structured tables
        metadata={},
        pages=pages,
        used_ocr=True
    )


# =============================================================================
# MAIN PROCESSOR
# =============================================================================

def process_document(file_path: str | Path) -> ProcessedDocument:
    """
    Process any supported document and return structured output.
    Automatically detects document type and routes to the right processor.
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    doc_type = detect_document_type(file_path)
    print(f"Detected document type: {doc_type.value}")

    if doc_type == DocumentType.PDF_NATIVE:
        return process_native_pdf(file_path)
    elif doc_type == DocumentType.EXCEL:
        return process_excel(file_path)
    elif doc_type == DocumentType.CSV:
        return process_csv(file_path)
    elif doc_type == DocumentType.WORD:
        return process_word(file_path)
    elif doc_type in [DocumentType.IMAGE, DocumentType.PDF_SCANNED]:
        return process_with_ocr(file_path, doc_type)
    else:
        raise ValueError(f"Unsupported document type: {file_path.suffix}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python document_processor.py <file_path> [--json]")
        print("\nSupported formats:")
        print("  - Images: .jpg, .jpeg, .png, .tiff, .bmp, .gif, .webp")
        print("  - PDF: .pdf (auto-detects native vs scanned)")
        print("  - Excel: .xlsx, .xls")
        print("  - CSV: .csv")
        print("  - Word: .docx")
        sys.exit(1)

    file_path = sys.argv[1]
    output_json = "--json" in sys.argv

    result = process_document(file_path)

    if output_json:
        print(json.dumps(asdict(result), indent=2, default=str))
    else:
        print(f"\n{'='*60}")
        print(f"Source: {result.source_file}")
        print(f"Type: {result.document_type}")
        print(f"Pages: {result.pages}")
        print(f"Used OCR: {result.used_ocr}")
        print(f"{'='*60}")
        print("\nTEXT CONTENT:")
        print("-" * 40)
        print(result.text_content[:2000])  # First 2000 chars
        if len(result.text_content) > 2000:
            print(f"\n... [{len(result.text_content) - 2000} more characters]")
        print("-" * 40)

        if result.tables:
            print(f"\nTABLES FOUND: {len(result.tables)}")


if __name__ == "__main__":
    main()
